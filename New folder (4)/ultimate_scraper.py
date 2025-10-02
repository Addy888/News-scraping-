
import os
import re
import time
import hashlib
import random
import urllib3
from pathlib import Path
from datetime import datetime
from collections import Counter
from io import BytesIO

import streamlit as st
import requests
from bs4 import BeautifulSoup
from langdetect import detect, LangDetectException

# optional cloudscraper
try:
    import cloudscraper
except Exception:
    cloudscraper = None

# selenium imports
from selenium import webdriver
from selenium.common.exceptions import WebDriverException, SessionNotCreatedException
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

# outputs
from docx import Document
from fpdf import FPDF
from PIL import Image

# ---------------- Config ----------------
OUTPUT_BASE = "outputs_final"
os.makedirs(OUTPUT_BASE, exist_ok=True)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)  # when verify=False used

# attempt auto-download NotoSans for Unicode PDF if missing
FONT_NAME = "NotoSans-Regular.ttf"
FONT_URL = "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf"
if not os.path.exists(FONT_NAME):
    try:
        r = requests.get(FONT_URL, timeout=20)
        if r.status_code == 200:
            with open(FONT_NAME, "wb") as f:
                f.write(r.content)
    except Exception:
        pass

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/140.0.7339.210 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:139.0) Gecko/20100101 Firefox/139.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15"
]

# ---------------- Helpers ----------------
def short_safe_folder(title: str, base_dir: str = OUTPUT_BASE, max_len=60) -> str:
    if not title:
        title = "article"
    clean = re.sub(r'[^A-Za-z0-9_]', '_', title)
    clean = re.sub(r'_{2,}', '_', clean).strip('_')
    if len(clean) <= max_len:
        short = clean
    else:
        h = hashlib.md5(clean.encode('utf-8')).hexdigest()[:8]
        short = clean[: max_len - 9] + "_" + h
    if not short:
        short = "article_" + datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(base_dir, short)

def detect_language(text: str) -> str:
    try:
        return detect(text[:1000])
    except LangDetectException:
        return "hi" if re.search(r'[\u0900-\u097F]', text) else "en"
    except Exception:
        return "en"

def split_sentences(text: str, lang='en'):
    text = text.strip()
    if not text:
        return []
    if lang.startswith('hi') or re.search(r'[\u0900-\u097F]', text):
        sents = re.split(r'(?<=[।!?])\s+|(?<=[!?])\s+|\n+', text)
    else:
        sents = re.split(r'(?<=[.!?])\s+|\n+', text)
    return [s.strip() for s in sents if s.strip()]

def extractive_summary(text: str, max_sentences: int = 50, lang: str = 'en') -> str:
    sents = split_sentences(text, lang)
    if not sents:
        return ""
    if len(sents) <= max_sentences:
        return " ".join(sents)
    words = [w.lower() for w in re.findall(r'\w+', text)]
    freq = Counter(words)
    scores = {i: sum(freq.get(w.lower(), 0) for w in re.findall(r'\w+', s)) for i, s in enumerate(sents)}
    top_idx = sorted(scores, key=scores.get, reverse=True)[:max_sentences]
    top_sorted = sorted(top_idx)
    return " ".join([sents[i] for i in top_sorted])

def paraphrase_text(text: str, lang: str = 'en') -> str:
    if not text:
        return ""
    if lang.startswith('hi') or re.search(r'[\u0900-\u097F]', text):
        syn = {'कहा':'बताया','बताया':'सूचित किया','किया':'अंजाम दिया'}
        out = text
        for k,v in syn.items():
            out = out.replace(k,v)
        return out
    else:
        syn = {'said':'stated','shows':'reveals','important':'crucial','use':'utilize','many':'numerous'}
        out = text
        for k,v in syn.items():
            out = re.sub(r'\b' + re.escape(k) + r'\b', v, out, flags=re.IGNORECASE)
        return out

# ---------------- Fetch strategies ----------------
def fetch_requests(url: str, verify_ssl: bool = True, timeout: int = 15) -> str:
    headers = {
        "User-Agent": random.choice(USER_AGENTS),
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": "https://www.google.com"
    }
    r = requests.get(url, headers=headers, timeout=timeout, verify=verify_ssl)
    r.raise_for_status()
    return r.text

def fetch_cloudscraper(url: str, verify_ssl: bool = True, timeout: int = 20) -> str:
    if cloudscraper is None:
        raise RuntimeError("cloudscraper not installed")
    scraper = cloudscraper.create_scraper()
    if not verify_ssl:
        try:
            scraper.session.verify = False
        except Exception:
            pass
    r = scraper.get(url, timeout=timeout)
    r.raise_for_status()
    return r.text

def fetch_selenium(url: str, headless: bool = True, wait_time: int = 5, scroll: bool = True, max_wait: int = 15) -> str:
    options = webdriver.ChromeOptions()
    # important flag to avoid invalid argument on many recent Chrome/driver combos
    options.add_argument("--remote-allow-origins=*")
    if headless:
        try:
            options.add_argument("--headless=new")
        except Exception:
            options.add_argument("--headless")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-blink-features=AutomationControlled")
    # random UA + window size
    ua = random.choice(USER_AGENTS)
    options.add_argument(f"--user-agent={ua}")
    w = random.choice([1200,1366,1440,1600])
    h = random.choice([700,768,800,900])
    options.add_argument(f"--window-size={w},{h}")
    options.add_experimental_option("excludeSwitches", ["enable-automation", "enable-logging"])
    options.add_experimental_option("useAutomationExtension", False)

    try:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
    except SessionNotCreatedException as sce:
        raise RuntimeError(f"SessionNotCreatedException: {sce}")
    except WebDriverException as we:
        raise RuntimeError(f"WebDriverException: {we}")
    except Exception as e:
        raise RuntimeError(f"Failed to start Chrome WebDriver: {e}")

    try:
        driver.get(url)
        try:
            WebDriverWait(driver, max_wait).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        except Exception:
            pass
        time.sleep(wait_time)
        if scroll:
            for _ in range(random.randint(1, 3)):
                driver.execute_script("window.scrollBy(0, window.innerHeight * 0.7);")
                time.sleep(random.uniform(0.6, 1.4))
            driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(0.3)
        return driver.page_source
    finally:
        try:
            driver.quit()
        except Exception:
            pass

# ---------------- Extraction ----------------
def extract_text_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "iframe"]):
        tag.decompose()
    main = soup.find("main") or soup.find("article")
    if main:
        ps = [p.get_text(" ", strip=True) for p in main.find_all("p") if p.get_text(strip=True)]
        if ps:
            return "\n\n".join(ps)
    ps = [p.get_text(" ", strip=True) for p in soup.find_all("p") if p.get_text(strip=True)]
    return "\n\n".join(ps)

def extract_title_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    m = soup.find("meta", property="og:title")
    if m and m.get("content"):
        return m.get("content").strip()
    return ""

# ---------------- Images (page-only) ----------------
def download_images_from_html(html: str, folder: str, ignore_ssl: bool = False, max_images: int = 200) -> list:
    soup = BeautifulSoup(html, "html.parser")
    imgs = soup.find_all("img")
    saved_paths = []
    headers = {"User-Agent": random.choice(USER_AGENTS), "Referer": "https://www.google.com"}
    c = 0
    for img in imgs:
        if c >= max_images:
            break
        src = img.get("src") or img.get("data-src") or img.get("data-lazy-src")
        if not src:
            continue
        src = src.strip()
        if not src.lower().startswith("http"):
            continue
        try:
            r = requests.get(src, headers=headers, timeout=20, verify=(not ignore_ssl), stream=True)
            r.raise_for_status()
            ext = Path(src).suffix if Path(src).suffix and len(Path(src).suffix) <= 6 else ".jpg"
            filename = f"img_{c+1}{ext}"
            outpath = os.path.join(folder, filename)
            with open(outpath, "wb") as f:
                for chunk in r.iter_content(8192):
                    if chunk:
                        f.write(chunk)
            # verify image
            try:
                Image.open(outpath).verify()
            except Exception:
                try: os.remove(outpath)
                except: pass
                continue
            saved_paths.append(outpath)
            c += 1
        except Exception:
            continue
    return saved_paths

# ---------------- Save outputs ----------------
def save_docx(title: str, url: str, summary: str, filepath: str):
    d = Document()
    d.add_heading(title or "Summary", level=1)
    d.add_paragraph(f"URL: {url}\n\n")
    d.add_paragraph(summary)
    d.save(filepath)

def save_pdf(title: str, url: str, summary: str, filepath: str):
    pdf = FPDF()
    pdf.add_page()
    if os.path.exists(FONT_NAME):
        try:
            pdf.add_font("NotoSans", "", FONT_NAME, uni=True)
            pdf.set_font("NotoSans", size=12)
        except Exception:
            pdf.set_font("Arial", size=12)
    else:
        pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 6, f"{title}\n\nURL: {url}\n\n{summary}")
    pdf.output(filepath)

def save_txt(title: str, url: str, summary: str, filepath: str):
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(f"Title: {title}\nURL: {url}\n\nSummary:\n{summary}")

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="Final Scraper", layout="wide")
st.title("🚀 Final Scraper — Title/Text/Images → DOCX/PDF/TXT")

st.markdown("**Caution:** This tool can optionally bypass SSL and anti-bot protections. Use only on sites you own or have permission to scrape.")

left, right = st.columns([3,1])
with left:
    mode = st.radio("Mode", ["Single", "Batch"])
    if mode == "Single":
        url_input = st.text_input("Enter URL (only this URL will be scraped)")
    else:
        url_input = st.text_area("Enter URLs (one per line). Only these will be scraped.")
with right:
    headless = st.checkbox("Run browser headless", value=True)
    wait_time = st.slider("JS wait (seconds)", 2, 12, 5)
    summary_sentences = st.slider("Summary sentences", 1, 200, 50)
    do_paraphrase = st.checkbox("Paraphrase (offline)", value=True)
    ignore_ssl = st.checkbox("Ignore SSL certificate errors for downloads (insecure)", value=True)
    max_images = st.number_input("Max images per page", min_value=1, max_value=500, value=50)

st.markdown("---")
agree = st.checkbox("I accept the risk and enable bypass (insecure)", value=False)
if agree:
    st.warning("Bypass enabled — script may ignore SSL for downloads and use cloudscraper/Selenium to fetch content.")

# pipeline
def process_url_pipeline(url: str):
    url = url.strip()
    if not url:
        return None, "empty"
    html = None
    # 1) try requests
    try:
        html = fetch_requests(url, verify_ssl=(not ignore_ssl))
    except Exception:
        html = None
    # 2) try cloudscraper if available and bypass enabled or requests failed
    if html is None and cloudscraper is not None:
        try:
            html = fetch_cloudscraper(url, verify_ssl=(not ignore_ssl))
        except Exception:
            html = None
    # 3) selenium fallback
    if html is None or ("access denied" in (html or "").lower()) or ("Request blocked" in (html or "")):
        try:
            html = fetch_selenium(url, headless=headless, wait_time=wait_time, scroll=True, max_wait=20)
        except Exception as e:
            return None, f"fetch_error: {e}"

    if not html:
        return None, "no_html"

    low = html.lower()
    if "access denied" in low and "reference" in low:
        return None, "access_denied"

    text = extract_text_from_html(html)
    lang = detect_language(text or "")
    summary = extractive_summary(text or "", max_sentences=summary_sentences, lang=lang)
    paraphrased = paraphrase_text(summary, lang=lang) if do_paraphrase else summary
    title = extract_title_from_html(html) or url.split("/")[-1] or "article"
    folder = short_safe_folder(title)
    try:
        os.makedirs(folder, exist_ok=True)
    except Exception:
        folder = os.path.join(OUTPUT_BASE, "article_" + hashlib.md5(title.encode('utf-8')).hexdigest()[:8])
        os.makedirs(folder, exist_ok=True)

    images = download_images_from_html(html, folder, ignore_ssl=ignore_ssl, max_images=max_images)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(folder, f"summary_{ts}.docx")
    pdf_path = os.path.join(folder, f"summary_{ts}.pdf")
    txt_path = os.path.join(folder, f"summary_{ts}.txt")
    try:
        save_docx(title, url, paraphrased, docx_path)
        save_pdf(title, url, paraphrased, pdf_path)
        save_txt(title, url, paraphrased, txt_path)
    except Exception as e:
        return None, f"save_error: {e}"

    return {
        "url": url,
        "title": title,
        "folder": folder,
        "images": images,
        "docx": docx_path,
        "pdf": pdf_path,
        "txt": txt_path,
        "summary": paraphrased
    }, "ok"

# UI actions
results = []
if mode == "Single":
    if st.button("Scrape & Save"):
        if not url_input.strip():
            st.warning("Enter a URL")
        else:
            if not agree:
                st.info("Bypass not enabled — cloudscraper and SSL-ignore will not be used.")
            with st.spinner("Working — may take a few seconds..."):
                res, status = process_url_pipeline(url_input)
            if res and status == "ok":
                st.success("Saved.")
                st.write("Title:", res["title"])
                st.write("Summary:", res["summary"])
                st.write("Folder:", res["folder"])
                if res["images"]:
                    st.write(f"Images downloaded: {len(res['images'])}")
                    for p in res["images"][:50]:
                        st.write(p)
                try:
                    with open(res["txt"], "rb") as fh:
                        st.download_button("Download TXT", data=fh.read(), file_name=os.path.basename(res["txt"]))
                except Exception:
                    pass
            else:
                st.error(f"Failed: {status}")
else:
    if st.button("Run Batch"):
        lines = [u.strip() for u in url_input.splitlines() if u.strip()]
        if not lines:
            st.warning("Paste URLs (one per line)")
        else:
            succ = 0
            fail = 0
            with st.spinner(f"Processing {len(lines)} URLs — this may take time..."):
                for u in lines:
                    res, status = process_url_pipeline(u)
                    if res and status == "ok":
                        succ += 1
                        results.append(res)
                    else:
                        fail += 1
            st.success(f"Batch complete — Success: {succ}, Failures: {fail}")
            if results:
                for r in results[:10]:
                    st.write(r["title"], "-", r["folder"], "-", len(r["images"]), "images")

st.markdown("---")
st.caption("Tool scrapes ONLY the URLs you provide. Bypass mode is insecure. Use responsibly and only with permission.")
